import json
import os
import shutil
import signal
import time
from datetime import datetime
from zoneinfo import ZoneInfo


def _patch_gradio_client():
    try:
        from gradio_client import utils as gc_utils
    except ImportError:
        return

    if hasattr(gc_utils, "get_type"):
        _orig_gt = gc_utils.get_type

        def _safe_get_type(schema):
            if not isinstance(schema, dict):
                return "Any"
            try:
                return _orig_gt(schema)
            except Exception:
                return "Any"

        gc_utils.get_type = _safe_get_type

    if hasattr(gc_utils, "_json_schema_to_python_type"):
        _orig_j = gc_utils._json_schema_to_python_type

        def _safe_json(schema, defs=None):
            if not isinstance(schema, dict):
                return "Any"
            try:
                return _orig_j(schema, defs)
            except Exception:
                return "Any"

        gc_utils._json_schema_to_python_type = _safe_json

    if hasattr(gc_utils, "json_schema_to_python_type"):
        _orig_j2 = gc_utils.json_schema_to_python_type

        def _safe_json2(schema):
            try:
                return _orig_j2(schema)
            except Exception:
                return "Any"

        gc_utils.json_schema_to_python_type = _safe_json2


_patch_gradio_client()

import gradio as gr
from google import genai
from google.genai import types
from PIL import Image
import webbrowser
import threading
import time

# =============== [ PATH / TIME ] ===============
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
HISTORY_DIR = os.path.join(BASE_DIR, "history")
HISTORY_IMG_DIR = os.path.join(HISTORY_DIR, "images")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
ENV_PATH = os.path.join(BASE_DIR, ".env")

KST = ZoneInfo("Asia/Seoul")

def _now_kst() -> datetime:
    return datetime.now(KST)

def _ts_file() -> str:
    """파일명용 타임스탬프 (KST, 초 단위까지 → 고유성 보장)."""
    return _now_kst().strftime("%Y%m%d_%H%M%S")

def _ts_session() -> str:
    """세션/프로젝트명용 타임스탬프 (KST, 분 단위까지)."""
    return _now_kst().strftime("%Y%m%d_%H%M")

def _extract_text_from_doc(filepath: str) -> str:
    """지원되지 않는 문서(Word, Excel 등)의 텍스트를 로컬에서 직접 추출합니다."""
    ext = filepath.lower().rsplit(".", 1)[-1]
    text = ""
    try:
        if ext == "docx":
            from docx import Document
            doc = Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif ext in ("xlsx", "xls"):
            import pandas as pd
            df = pd.read_excel(filepath)
            text = df.to_string(index=False)
        elif ext == "csv":
            import pandas as pd
            df = pd.read_csv(filepath)
            text = df.to_string(index=False)
        elif ext == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
    except Exception:
        pass
    return text.strip()


# =============== [ API KEY / CLIENT ] ===============
_CLIENTS: dict = {}

def get_client(api_key: str) -> genai.Client:
    if api_key not in _CLIENTS:
        _CLIENTS[api_key] = genai.Client(api_key=api_key)
    return _CLIENTS[api_key]

def load_api_key() -> str:
    if not os.path.exists(ENV_PATH):
        return ""
    try:
        with open(ENV_PATH, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip().startswith("GEMINI_API_KEY="):
                    return line.strip().split("=", 1)[1]
    except Exception:
        pass
    return ""

def save_api_key(api_key: str) -> None:
    if not api_key or api_key == load_api_key():
        return
    with open(ENV_PATH, "w", encoding="utf-8") as f:
        f.write(f"GEMINI_API_KEY={api_key}\n")


# =============== [ IMAGE STUDIO ] ===============
IMAGE_MODEL_MAPPING = {
    "Gemini 3.1 Flash Image (1K: ~$0.067/장)": "gemini-3.1-flash-image-preview",
    "Gemini 3 Pro Image (1K/2K: ~$0.134/장)": "gemini-3-pro-image-preview",
    "Gemini 2.5 Flash Image (1K: ~$0.039/장)": "gemini-2.5-flash-image",
    "Imagen 4 Standard ($0.04/장)": "imagen-4.0-generate-001",
    "Imagen 4 Ultra ($0.06/장)": "imagen-4.0-ultra-generate-001",
    "Imagen 4 Fast ($0.02/장)": "imagen-4.0-fast-generate-001",
}

ASPECT_RATIOS = ["1:1", "1:4", "1:8", "2:3", "3:2", "3:4", "4:1", "4:3", "4:5", "5:4", "8:1", "9:16", "16:9", "21:9"]
RESOLUTIONS = ["512", "1K", "2K", "4K"]
THINK_LEVELS = ["Minimal", "High"]
IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "heic"}

SAFETY_SETTINGS = [
    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
]

SYSTEM_INSTRUCTION = (
    "당신은 최고 수준의 시각적 요소를 전문으로 생성하고 다루는 전문가입니다. "
    "완벽한 디자인과 시각 수정 결과를 배출하세요."
)


def get_image_config(aspect_ratio, thinking_level, model_id):
    img_args = {"aspect_ratio": aspect_ratio} if aspect_ratio else {}
    
    # 2.5-flash-image 모델은 thinking과 google_search 도구를 지원하지 않음
    is_gemini_3 = "gemini-3" in model_id.lower()
    
    think_config = None
    if thinking_level and thinking_level.lower() == "high" and is_gemini_3:
        think_config = types.ThinkingConfig(thinking_budget=-1, include_thoughts=True)

    tools_list = None
    if is_gemini_3:
        tools_list = [types.Tool(google_search=types.GoogleSearch())]

    # 2.5-flash-image의 경우 확실하게 이미지만 강제 (텍스트 회피)
    modalities = ["TEXT", "IMAGE"] if is_gemini_3 else ["IMAGE"]

    return types.GenerateContentConfig(
        response_modalities=modalities,
        image_config=types.ImageConfig(**img_args) if img_args else None,
        thinking_config=think_config,
        tools=tools_list,
        system_instruction=SYSTEM_INSTRUCTION,
        safety_settings=SAFETY_SETTINGS,
    )


def _build_final_prompt(history, prompt, is_imagen, resolution):
    """직전 대화 맥락을 병합해 최종 프롬프트를 만든다."""
    base = prompt.strip() if prompt and prompt.strip() else "수정/재생성 진행"

    if is_imagen:
        prev = [u for u, _ in history if isinstance(u, str)]
        if prev:
            base = " ".join(prev) + " " + base
    else:
        lines = []
        for u, m in history:
            if isinstance(u, str):
                lines.append(f"User: {u}")
            elif isinstance(u, (tuple, list)):
                lines.append("[사용자가 이전 턴에 이미지를 포함했습니다]")
            if isinstance(m, str):
                lines.append(f"AI: {m}")
        if lines:
            base = "이전 대화 맥락:\n" + "\n".join(lines) + f"\n\n이번 생성 지시사항:\n{base}"

    if resolution:
        base += f" \n[지시: {resolution} 해상도 타겟]"
    return base


def _generate_session_title(client, prompt):
    if not prompt:
        return "프로젝트"
    try:
        r = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=(
                "다음 프롬프트를 바탕으로 프로젝트 폴더명을 15자 이내의 핵심적인 단어들로만 만들어줘 "
                f"(예: 파란머리_사이버펑크). 특수문자 절대 금지, 띄어쓰기는 _로: {prompt}"
            ),
        )
        title = r.text.strip().replace(" ", "_").replace("\n", "").replace('"', "").replace("'", "")[:15]
        return title or "프로젝트"
    except Exception:
        return "프로젝트"


def _run_imagen(client, model_id, prompt, aspect_ratio, num_imgs, out_prefix):
    cfg_kwargs = {"output_mime_type": "image/png", "number_of_images": num_imgs}
    if aspect_ratio:
        cfg_kwargs["aspect_ratio"] = aspect_ratio
    res = client.models.generate_images(
        model=model_id, prompt=prompt, config=types.GenerateImagesConfig(**cfg_kwargs)
    )
    ts = _ts_file()
    paths = []
    for idx, g in enumerate(res.generated_images):
        path = os.path.join(OUTPUT_DIR, f"{out_prefix}_{ts}_{idx}.png")
        g.image.save(path)
        paths.append(path)
    return paths


def _save_history_json(history, session_id):
    """대화 히스토리를 history/proj_{session_id}.json으로 저장 (이미지는 images/ 폴더로 복제)."""
    os.makedirs(HISTORY_IMG_DIR, exist_ok=True)

    def _proc(msg):
        if isinstance(msg, dict) and "file" in msg:
            path = msg["file"]
        elif isinstance(msg, (tuple, list)) and len(msg) > 0:
            path = msg[0]
        else:
            return msg
        if not (isinstance(path, str) and os.path.exists(path)):
            return msg
        dname = os.path.basename(os.path.dirname(path))
        fname = os.path.basename(path) if dname in ("outputs", "images") else f"{dname}_{os.path.basename(path)}"
        new_path = os.path.join(HISTORY_IMG_DIR, fname)
        if not os.path.exists(new_path) and os.path.abspath(path) != os.path.abspath(new_path):
            shutil.copy2(path, new_path)
        return (new_path,)

    clean = [(_proc(u), _proc(b)) for u, b in history]
    with open(os.path.join(HISTORY_DIR, f"proj_{session_id}.json"), "w", encoding="utf-8") as f:
        json.dump({"session": session_id, "history": clean}, f, ensure_ascii=False, indent=2)


def process_image_interaction(api_key, model_display_name, prompt, upload_files, history, session_id,
                              aspect_ratio, resolution, thinking_level, num_images_img):
    if not session_id:
        session_id = f"NEW_{_ts_file()}"

    if not api_key:
        history.append((prompt, "❌ API 키가 입력되어야 합니다."))
        yield history, session_id, "", None
        return

    save_api_key(api_key)
    os.environ["GEMINI_API_KEY"] = api_key

    if upload_files:
        for f in upload_files:
            history.append(((f,), None))
    if prompt:
        history.append((prompt, None))

    history.append((None, "✨ Gemini가 상상력을 동원하는 중..."))
    yield history, session_id, "", None

    try:
        client = get_client(api_key)

        # 새 세션이면 AI로 프로젝트명을 만들어 KST 연월일시분과 결합
        if session_id.startswith("NEW_"):
            title = _generate_session_title(client, prompt)
            session_id = f"{_ts_session()}_{title}"

        model_id = IMAGE_MODEL_MAPPING.get(model_display_name, "gemini-3.1-flash-image-preview")
        is_imagen = "imagen" in model_id.lower()

        final_prompt = _build_final_prompt(history[:-1], prompt, is_imagen, resolution)

        history.pop()  # 로딩 메시지 제거
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        parsed_docs = []
        gemini_target_files = []
        if upload_files:
            for f in upload_files:
                ext = f.lower().rsplit(".", 1)[-1]
                if ext in IMAGE_EXTS:
                    gemini_target_files.append(Image.open(f))
                else:
                    history.append((None, f"✨ 문서를 분석/전송 중... ({os.path.basename(f)[:10]})"))
                    yield history, session_id, gr.update(), gr.update()
                    try:
                        ext_lower = f.lower().rsplit(".", 1)[-1]
                        if ext_lower in ("docx", "xlsx", "xls", "csv", "txt"):
                            doc_text = _extract_text_from_doc(f)
                            if doc_text:
                                if is_imagen:
                                    parsed_docs.append(f"[{os.path.basename(f)} 내용]\n{doc_text[:10000]}")
                                else:
                                    gemini_target_files.append(f"[{os.path.basename(f)} 내용]\n{doc_text[:10000]}")
                                history.pop()
                                continue

                        api_f = client.files.upload(file=f)
                        if is_imagen:
                            parsed_docs.append(api_f)
                        else:
                            gemini_target_files.append(api_f)
                    except Exception as upload_err:
                        print(f"문서 업로드 실패: {upload_err}")
                    history.pop()

        if is_imagen and parsed_docs:
            history.append((None, "✨ Imagen 모델을 위해 문서 컨텍스트 프롬프트로 변환 중..."))
            yield history, session_id, gr.update(), gr.update()
            try:
                doc_res = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=parsed_docs + [
                        "이 첨부된 문서(들)의 내용과 특징을 분석하여, 이 문서들을 완벽하게 대변할 수 있는 "
                        "고품질의 사진/그림 생성용 묘사(프롬프트)로 구성해줘. 한국어로 묘사만 출력해."
                    ],
                )
                final_prompt = f"[첨부 문서 분석 결과 묘사: {doc_res.text}]\n\n{final_prompt}"
            except Exception:
                pass
            history.pop()

        output_images = []
        output_texts = []
        num_imgs = int(num_images_img) if num_images_img else 1

        if is_imagen:
            output_images = _run_imagen(client, model_id, final_prompt, aspect_ratio, num_imgs, "img")
            output_texts.append("✨ Imagen 렌더링 완료")
        else:
            config = get_image_config(aspect_ratio, thinking_level, model_id)
            for req_i in range(num_imgs):
                contents = [final_prompt] if final_prompt else []
                contents.extend(gemini_target_files)
                res = client.models.generate_content(model=model_id, contents=contents, config=config)

                grounding = getattr(res.candidates[0], "grounding_metadata", None) if res.candidates else None
                if (grounding and getattr(grounding, "search_entry_point", None)
                        and "🔍 Google Search 그라운딩 반영됨" not in output_texts):
                    output_texts.append("🔍 Google Search 그라운딩 반영됨")

                ts = _ts_file()
                for part in res.parts:
                    if getattr(part, "thought", False):
                        if part.text and req_i == 0:
                            output_texts.append(f"💭 [사고]: {part.text.replace('**', '')}")
                        continue
                    if part.text and req_i == 0:
                        output_texts.append(part.text.replace("**", ""))
                    elif hasattr(part, "as_image"):
                        try:
                            img = part.as_image()
                            if img:
                                path = os.path.join(OUTPUT_DIR, f"img_{ts}_{req_i}.png")
                                img.save(path)
                                output_images.append(path)
                        except Exception:
                            pass

        for path in output_images:
            history.append((None, (path,)))
        if output_texts:
            history.append((None, "\n\n".join(output_texts)))

    except Exception as e:
        if history and history[-1][1] and "✨ Gemini" in str(history[-1][1]):
            history.pop()
        history.append((None, f"🚨 로컬/서버 오류: {str(e).replace('**', '')}"))

    try:
        _save_history_json(history, session_id)
    except Exception:
        pass

    yield history, session_id, gr.update(value=""), gr.update(value=None)


def reset_session():
    return [], ""


def generate_modal_image(api_key, model_display_name, new_prompt, prev_prompt_str, modal_img_path,
                         history, row_idx_str, resolution):
    yield gr.update(value="⏳ 처리 시작..."), gr.update(), gr.update()

    if not api_key:
        yield gr.update(value="❌ API 키가 필요합니다."), gr.update(), gr.update()
        return

    try:
        client = get_client(api_key)
        model_id = IMAGE_MODEL_MAPPING.get(model_display_name, "gemini-3.1-flash-image-preview")
        is_imagen = "imagen" in model_id.lower()

        row_idx = int(row_idx_str) if row_idx_str else 0
        sliced_history = history[:row_idx]

        effective_prompt = new_prompt.strip() if new_prompt and new_prompt.strip() else prev_prompt_str
        final_prompt = _build_final_prompt(sliced_history, effective_prompt, is_imagen, resolution)

        yield gr.update(value="✨ 상상력 동원 중..."), gr.update(), gr.update()
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        new_img_out = None
        log_text = ""

        if is_imagen:
            paths = _run_imagen(client, model_id, final_prompt, "1:1", 1, "img_modal")
            if paths:
                new_img_out = paths[0]
                log_text = "✨ Imagen 재생성 완료"
        else:
            contents = [final_prompt]
            if modal_img_path and os.path.exists(modal_img_path):
                contents.append(Image.open(modal_img_path))
            config = get_image_config("1:1", "Minimal", model_id)
            res = client.models.generate_content(model=model_id, contents=contents, config=config)
            ts = _ts_file()
            for part in res.parts:
                if getattr(part, "thought", False):
                    continue
                if hasattr(part, "as_image"):
                    try:
                        img = part.as_image()
                        if img:
                            path = os.path.join(OUTPUT_DIR, f"img_modal_{ts}.png")
                            img.save(path)
                            new_img_out = path
                    except Exception:
                        pass
                if part.text:
                    log_text += part.text + "\n"

        if new_img_out:
            replacement_val = (effective_prompt, (new_img_out,))
            yield gr.update(value=f"✅ 완료!\n{log_text}".strip()), gr.update(value=new_img_out), replacement_val
        else:
            yield gr.update(value="🚨 생성된 이미지가 없습니다."), gr.update(), gr.update()

    except Exception as e:
        yield gr.update(value=f"🚨 로컬/서버 오류: {str(e).replace('**', '')}"), gr.update(), gr.update()


def handle_modal_close(history, row_idx_str, pending_replacement):
    if pending_replacement and row_idx_str:
        try:
            row = int(row_idx_str)
            if row < len(history):
                user_p, ai_r = pending_replacement
                old_u, _ = history[row]
                if row > 0 and history[row - 1][1] is None:
                    history[row - 1] = (user_p, None)
                    history[row] = (old_u, ai_r)
                else:
                    history[row] = (user_p if user_p else old_u, ai_r)
        except Exception:
            pass
    return history, gr.update(visible=False), None


# =============== [ VIDEO STUDIO (Veo) ] ===============
VIDEO_MODEL_MAPPING = {
    "Veo 3.1 Standard (1080p: $0.40/초, 4k: $0.60/초)": "veo-3.1-generate-preview",
    "Veo 3.1 Fast (1080p: $0.12/초, 4k: $0.30/초)": "veo-3.1-fast-generate-preview",
    "Veo 3.1 Lite (1080p: $0.08/초, 4k 불가)": "veo-3.1-lite-generate-preview",
    "Veo 3.0 Standard ($0.40/초)": "veo-3.0-generate-001",
    "Veo 3.0 Fast (1080p: $0.12/초, 4k: $0.30/초)": "veo-3.0-fast-generate-001",
}


def _generate_video_core(api_key, model_display, prompt, first_fr, last_fr, refs,
                         aspect, res, dur, auto_extend, target_dur, state_obj, extend_last):
    if not api_key:
        yield "❌ API 키가 필요합니다.", None, state_obj
        return

    save_api_key(api_key)
    os.environ["GEMINI_API_KEY"] = api_key
    yield "⏳ Veo 동영상 렌더링 서버 접속 중...", None, state_obj

    try:
        client = genai.Client()
        model_id = VIDEO_MODEL_MAPPING.get(model_display, "veo-3.1-generate-preview")
        kwargs = {"model": model_id}
        if prompt:
            kwargs["prompt"] = prompt.replace("**", "")

        if extend_last:
            if not state_obj:
                yield "❌ 연장할 기존 영상 캐시(서버 API 객체)가 없습니다. 먼저 새 영상을 생성해주세요.", None, state_obj
                return
            kwargs["video"] = state_obj
            res = "720p"

        if first_fr and not extend_last:
            with open(first_fr, "rb") as f_img:
                kwargs["image"] = types.Image(image_bytes=f_img.read(), mime_type="image/png")

        config_args = {}
        if aspect and not extend_last:
            config_args["aspect_ratio"] = aspect
        if last_fr and not extend_last:
            with open(last_fr, "rb") as l_img:
                # Some models do not support last_frame, but if passed it must be types.Image
                config_args["last_frame"] = types.Image(image_bytes=l_img.read(), mime_type="image/png")

        if refs and not extend_last:
            ref_objs = []
            parsed_docs = []
            for r in refs:
                ext = r.lower().rsplit(".", 1)[-1]
                if ext in IMAGE_EXTS:
                    if len(ref_objs) < 3:
                        with open(r, "rb") as r_img:
                            ref_img_obj = types.Image(image_bytes=r_img.read(), mime_type="image/png")
                        ref_objs.append(types.VideoGenerationReferenceImage(
                            image=ref_img_obj, reference_type="asset"
                        ))
                else:
                    yield f"✨ 문서를 서버로 전송/분석 중... ({os.path.basename(r)[:10]})", None, state_obj
                    try:
                        if ext in ("docx", "xlsx", "xls", "csv", "txt"):
                            doc_text = _extract_text_from_doc(r)
                            if doc_text:
                                parsed_docs.append(f"[{os.path.basename(r)} 내용]\n{doc_text[:10000]}")
                                continue
                        parsed_docs.append(client.files.upload(file=r))
                    except Exception:
                        pass

            if ref_objs:
                config_args["reference_images"] = ref_objs
                dur = "8"

            if parsed_docs:
                yield "✨ 문서를 비디오 프롬프트로 변환 분석 중...", None, state_obj
                try:
                    doc_res = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=parsed_docs + [
                            "이 문서(들)를 철저히 분석해서 해당 내용을 완벽하게 대변하는 훌륭한 비디오 생성용 "
                            "시각적 프롬프트(장면 묘사, 분위기, 행동 등)로 변환해줘. "
                            "주의사항: 절대 불필요한 설명은 빼고 순수하게 한국어 '장면 묘사 프롬프트'만 출력할 것."
                        ],
                    )
                    prompt = f"[첨부 문서 기반 묘사: {doc_res.text}]\n\n{prompt or ''}"
                    kwargs["prompt"] = prompt.replace("**", "")
                except Exception:
                    pass

        if res in ("1080p", "4k"):
            dur = "8"
        if res:
            config_args["resolution"] = res
        if dur:
            config_args["duration_seconds"] = dur
        if config_args:
            kwargs["config"] = types.GenerateVideosConfig(**config_args)

        operation = client.models.generate_videos(**kwargs)
        current_dur = int(dur) if dur else 8
        target_duration = int(target_dur) if target_dur else current_dur

        while not operation.done:
            yield f"🎬 영상 렌더링 진행 중... (현재 {current_dur}초 분량 생성 중)", None, state_obj
            time.sleep(10)
            operation = client.operations.get(operation)

        native_vid_obj = operation.response.generated_videos[0].video

        while auto_extend and not extend_last and current_dur < target_duration:
            yield (f"🔄 목표 길이({target_duration}초) 도달을 위해 이어붙이기 파이프라인 가동... "
                   f"(현재 {current_dur}초 완료)"), None, native_vid_obj
            ext_kwargs = {
                "model": model_id,
                "video": native_vid_obj,
                "config": types.GenerateVideosConfig(person_generation="allow_all", resolution="720p"),
            }
            if prompt:
                ext_kwargs["prompt"] = prompt.replace("**", "")
            ext_op = client.models.generate_videos(**ext_kwargs)
            while not ext_op.done:
                yield f"🔄 {target_duration}초 스윙 연장 렌더링 중... (최장 5분 이상 소요 가능)", None, native_vid_obj
                time.sleep(10)
                ext_op = client.operations.get(ext_op)
            native_vid_obj = ext_op.response.generated_videos[0].video
            current_dur += 7

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        out_path = os.path.join(OUTPUT_DIR, f"veo_{_ts_file()}.mp4")
        try:
            client.files.download(file=native_vid_obj)
            native_vid_obj.save(out_path)
        except Exception:
            yield "⚠️ 파일 다운로드 구조 분석 중, 바이너리 저장 대체 방식을 시도합니다.", None, state_obj
            content = client.files.download(name=native_vid_obj.name)
            with open(out_path, "wb") as f:
                f.write(content)

        if os.path.exists(out_path):
            rel = os.path.relpath(out_path, BASE_DIR)
            yield f"✅ 렌더링 성공! 파일이 {rel} 에 저장되었습니다.", out_path, native_vid_obj
        else:
            yield "🚨 비디오 다운로드 검증 실패. 파일을 기록하지 못했습니다.", None, state_obj

    except Exception as e:
        yield f"🚨 비디오 생성/통신 파이프라인 오류: {str(e).replace('**', '')}", None, state_obj


def standard_vid_ui(*args):
    yield from _generate_video_core(*args, extend_last=False)

def extend_vid_ui(*args):
    yield from _generate_video_core(*args, extend_last=True)


# =============== [ UI HELPERS ] ===============
def _list_projects():
    if not os.path.exists(HISTORY_DIR):
        return gr.update(choices=[])
    projects = sorted(
        [f[:-5] for f in os.listdir(HISTORY_DIR) if f.endswith(".json")], reverse=True
    )
    return gr.update(choices=projects)


def _load_project(proj_name):
    if not proj_name:
        return [], ""
    fp = os.path.join(HISTORY_DIR, f"{proj_name}.json")
    if not os.path.exists(fp):
        return [], ""
    try:
        with open(fp, "r", encoding="utf-8") as f:
            data = json.load(f)
        h = []
        for u, b in data.get("history", []):
            if isinstance(u, list): u = tuple(u)
            if isinstance(b, list): b = tuple(b)
            h.append([u, b])
        return h, data.get("session", proj_name.replace("proj_", ""))
    except Exception:
        return [], ""


def _handle_image_select(evt: gr.SelectData, history):
    row, col = evt.index
    selected = history[row][col]
    filepath = None
    if isinstance(selected, tuple) and selected:
        filepath = selected[0]
    elif isinstance(selected, dict):
        filepath = selected.get("file")
    elif isinstance(selected, str) and selected.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
        filepath = selected

    if filepath and os.path.exists(filepath):
        prev_prompt = ""
        for r in range(row, -1, -1):
            cand = history[r][0]
            if isinstance(cand, str) and cand.strip():
                prev_prompt = cand
                break
        return gr.update(visible=True), filepath, filepath, str(row), prev_prompt, gr.update(value="대기 중..."), None
    return gr.update(visible=False), None, None, "", "", gr.update(), None


# =============== [ UI ] ===============
CSS = """
body.dark { font-family: 'Pretendard', sans-serif; background-color: #121212 !important; color: #ffffff !important; }
#chatbot-container, #video-container { border-radius: 12px; overflow: hidden; background-color: #1e1e1e !important; box-shadow: 0 4px 15px rgba(0,0,0,0.5); border: 1px solid #333;}
#chatbot-container img { max-height: 400px !important; max-width: 100% !important; object-fit: contain; }
#sidebar { background-color: #181818 !important; padding: 20px; border-radius: 12px; border: 1px solid #333; }
.modal-overlay {
    position: fixed !important; top: 0; left: 0; width: 100vw; height: 100vh; background: rgba(0,0,0,0.85) !important;
    z-index: 9999 !important; padding-top: 10vh;
}
.modal-box {
    background: #2a2a2a !important; padding: 24px; border-radius: 16px; min-width: 400px; max-width: 80%;
    margin: 0 auto !important; max-height: 90vh; overflow-y: auto;
    box-shadow: 0 10px 40px rgba(0,0,0,0.8) !important; z-index: 10000 !important; border: 1px solid #555 !important;
}
#header-actions { display: flex; justify-content: flex-end; align-items: flex-start; padding-top: 12px; }
#header-actions a, #header-actions button { white-space: nowrap; }
"""

JS_CODE = """
function() {
    document.body.classList.add("dark");
    document.querySelector("gradio-app").classList.add("dark");
    window.addEventListener("beforeunload", function (e) {
        navigator.sendBeacon('/shutdown');
    });
}
"""


def build_ui():
    with gr.Blocks(theme=gr.themes.Base(), css=CSS, js=JS_CODE) as demo:
        with gr.Row(elem_id="header-row"):
            with gr.Column(scale=8):
                gr.Markdown("# 🚀 Image&Video Laboratory")
                gr.Markdown("Made by Hyunho Kim, contact me hh09080@naver.com")
            with gr.Column(scale=1, min_width=140, elem_id="header-actions"):
                # 상대 URL 사용: 직접 실행(/manual) / 프록시 경유(/p/<port>/manual) 모두 정상 동작
                gr.Button("📖 사용설명서", link="manual", variant="secondary", size="sm")

        with gr.Tabs():
            # --- TAB 1: 이미지 스튜디오 ---
            with gr.Tab("🎨 이미지 스튜디오 (Gemini 3 Pro/Flash)"):
                session_id_state = gr.State("")
                with gr.Row():
                    with gr.Column(scale=3, elem_id="chatbot-container"):
                        chatbot = gr.Chatbot(height=600, type="tuples")
                        with gr.Row():
                            prompt = gr.Textbox(
                                scale=8, show_label=False,
                                placeholder="채팅하듯 수정이나 생성을 지시하세요 (Enter: 전송 / Shift+Enter: 줄바꿈)",
                                container=False, lines=1, max_lines=15,
                            )
                            submit_btn = gr.Button("로켓 전송", variant="primary", scale=1)
                            stop_btn = gr.Button("⏹️ 중지", variant="stop", scale=1)
                        upload_files = gr.File(
                            file_count="multiple",
                            label="참조/수정 원본 사진 또는 각종 문서 (PDF, Word, TXT, Excel, PPT, 이미지 다중 업로드 가능)",
                        )

                    with gr.Column(scale=1, elem_id="sidebar"):
                        with gr.Row():
                            api_key_img = gr.Textbox(label="🔑 API Key", type="password", value=load_api_key(), scale=3)
                            gr.Button(
                                "🔑 API 발급 & 사용량 조회",
                                link="https://aistudio.google.com/usage?timeRange=last-28-days&project",
                                size="sm",
                            )
                        model_selector_img = gr.Dropdown(
                            choices=list(IMAGE_MODEL_MAPPING.keys()),
                            value="Gemini 3.1 Flash Image (1K: ~$0.067/장)", label="모델",
                        )
                        aspect_ratio_img = gr.Dropdown(choices=ASPECT_RATIOS, value="1:1", label="비율")
                        resolution_img = gr.Dropdown(choices=RESOLUTIONS, value="2K", label="해상도")
                        thinking_level_img = gr.Radio(choices=THINK_LEVELS, value="Minimal", label="사고 수준")
                        num_images_img = gr.Slider(minimum=1, maximum=10, step=1, value=1, label="연속 생성 횟수")
                        clear_btn = gr.Button("🗑️ 세션 초기화", variant="stop")

                        with gr.Accordion("📂 프로젝트 기록 불러오기", open=False):
                            project_loader = gr.Dropdown(choices=[], label="저장된 프로젝트 (히스토리)", interactive=True)
                            load_btn = gr.Button("프로젝트 열기")
                            project_loader.focus(fn=_list_projects, inputs=None, outputs=[project_loader], api_name=False)
                            load_btn.click(
                                fn=_load_project, inputs=[project_loader],
                                outputs=[chatbot, session_id_state], api_name=False,
                            )

                # --- Modal ---
                with gr.Column(visible=False, elem_classes="modal-overlay") as modal_ui:
                    with gr.Column(elem_classes="modal-box"):
                        gr.Markdown("### 🖼️ 이미지 통합 관리 (수정 및 재생성)")
                        modal_row_idx = gr.State("")
                        modal_prev_prompt = gr.State("")
                        modal_pending_replacement = gr.State(None)
                        with gr.Row(elem_classes="modal-body"):
                            with gr.Column(scale=1):
                                modal_img_display = gr.Image(
                                    label="출력 이미지 미리보기", interactive=False,
                                    type="filepath", height=130,
                                )
                            with gr.Column(scale=3):
                                modal_prompt = gr.Textbox(
                                    label="이 이미지를 기준(참조)으로 추가할 프롬프트(수정/재생성)",
                                    placeholder="예: 배경을 밤으로 바꿔줘...", lines=2,
                                )
                                modal_status = gr.Textbox(label="실행 로그", interactive=False, lines=2)
                                with gr.Row():
                                    modal_dl_btn = gr.DownloadButton("⬇️ 원본 파일 다운로드", variant="secondary")
                                    modal_edit_btn = gr.Button("✨ 프롬프트 반영 및 재생성", variant="primary")
                                modal_close_btn = gr.Button("✖️ 반영 후 패널 닫기 (기존 사진 대체)", variant="stop")

                img_inputs = [
                    api_key_img, model_selector_img, prompt, upload_files, chatbot, session_id_state,
                    aspect_ratio_img, resolution_img, thinking_level_img, num_images_img,
                ]
                img_outputs = [chatbot, session_id_state, prompt, upload_files]

                sub_evt = submit_btn.click(fn=process_image_interaction, inputs=img_inputs, outputs=img_outputs, api_name=False)
                prompt_evt = prompt.submit(fn=process_image_interaction, inputs=img_inputs, outputs=img_outputs, api_name=False)
                stop_btn.click(fn=None, inputs=None, outputs=None, cancels=[sub_evt, prompt_evt], api_name=False)
                clear_btn.click(fn=reset_session, outputs=[chatbot, session_id_state], api_name=False)

                chatbot.select(
                    fn=_handle_image_select, inputs=[chatbot],
                    outputs=[
                        modal_ui, modal_img_display, modal_dl_btn, modal_row_idx,
                        modal_prev_prompt, modal_status, modal_pending_replacement,
                    ],
                    api_name=False,
                )

                modal_close_btn.click(
                    fn=handle_modal_close,
                    inputs=[chatbot, modal_row_idx, modal_pending_replacement],
                    outputs=[chatbot, modal_ui, modal_pending_replacement], api_name=False,
                ).then(
                    fn=None, inputs=None, outputs=None,
                    js="() => { document.dispatchEvent(new KeyboardEvent('keydown', {key: 'Escape', bubbles: true})); }",
                )

                modal_edit_btn.click(
                    fn=generate_modal_image,
                    inputs=[
                        api_key_img, model_selector_img, modal_prompt, modal_prev_prompt,
                        modal_img_display, chatbot, modal_row_idx, resolution_img,
                    ],
                    outputs=[modal_status, modal_img_display, modal_pending_replacement], api_name=False,
                )

                def update_image_ui(model_name):
                    # 해상도
                    if "Imagen" in model_name:
                        res_up = gr.update(choices=["1K"], value="1K")
                    else:
                        res_up = gr.update(choices=RESOLUTIONS, value="2K")
                    
                    # 화면 비율
                    if "Imagen" in model_name:
                        aspect_up = gr.update(choices=["1:1", "3:4", "4:3", "16:9", "9:16"], value="1:1")
                    else:
                        aspect_up = gr.update(choices=ASPECT_RATIOS, value="1:1")
                        
                    # 사고 수준 (Thinking) - 2.5 Flash Image 및 Imagen 미지원
                    if "2.5 Flash Image" in model_name or "Imagen" in model_name:
                        think_up = gr.update(interactive=False, value="Minimal")
                    else:
                        think_up = gr.update(interactive=True)
                        
                    return res_up, aspect_up, think_up

                model_selector_img.change(
                    fn=update_image_ui, 
                    inputs=[model_selector_img], 
                    outputs=[resolution_img, aspect_ratio_img, thinking_level_img]
                )

            # --- TAB 2: 비디오 렌더링 ---
            with gr.Tab("🎬 비디오 렌더링 룸 (Veo 3.1)"):
                last_video_state = gr.State(None)
                with gr.Row():
                    with gr.Column(scale=3, elem_id="video-container"):
                        video_output = gr.Video(label="생성된 동영상 결과물", height=500)
                        video_status = gr.Textbox(label="시스템 상태 창", interactive=False, lines=2)
                        video_prompt = gr.Textbox(
                            label="💬 감독 메가폰 (프롬프트/음향 지시)", lines=3,
                            placeholder="예: '이거 진짜 멋지다!' 라고 남자가 소리칩니다. 네온사인 불빛이 반짝이는 사이버펑크 도시... (효과음/대사이해 지원)",
                        )
                        with gr.Row():
                            generate_vid_btn = gr.Button("🎥 처음부터 동영상 렌더링", variant="primary")
                            extend_vid_btn = gr.Button("🔄 방금 만든 영상 7초 이어서 추가 연장하기", variant="secondary")
                            stop_vid_btn = gr.Button("⏹️ 렌더링 강제 취소", variant="stop")

                    with gr.Column(scale=2, elem_id="sidebar"):
                        with gr.Row():
                            api_key_vid = gr.Textbox(label="🔑 API Key", type="password", value=load_api_key(), scale=3)
                            gr.Button(
                                "🔑 API 발급 & 사용량 조회",
                                link="https://aistudio.google.com/usage?timeRange=last-28-days&project",
                                size="sm",
                            )
                        video_model = gr.Dropdown(
                            choices=list(VIDEO_MODEL_MAPPING.keys()),
                            value="Veo 3.1 Standard (1080p: $0.40/초, 4k: $0.60/초)", label="모델",
                        )

                        with gr.Accordion("📐 세부 촬영/포맷 옵션", open=True):
                            v_aspect_ratio = gr.Dropdown(choices=["16:9", "9:16"], value="16:9", label="화면 비율 (유튜브/숏폼)")
                            v_resolution = gr.Dropdown(
                                choices=["720p", "1080p", "4k"], value="1080p",
                                label="화질 (1080p, 4k는 무조건 8초 렌더링)",
                            )
                            v_duration = gr.Dropdown(choices=["4", "6", "8"], value="8", label="기본 영상 길이 (초)")
                            gr.Markdown("---")
                            v_auto_extend = gr.Checkbox(label="자동으로 이어 붙여서 진행 (Auto-Extend)", value=False)
                            v_target_dur = gr.Slider(minimum=8, maximum=64, step=7, value=15, label="목표 총 길이 (초)", interactive=False)

                        with gr.Accordion("🖼️ 이미지 연계 지시사항 (문서/이미지-to-비디오)", open=False):
                            gr.Markdown("동영상의 시작/끝을 직접 결정하거나 일관성을 강제합니다. 문서를 올리면 비디오 프롬프트로 자동 변환됩니다.")
                            v_first_frame = gr.Image(type="filepath", label="첫 프레임 이미지 (시작 화면)")
                            v_last_frame = gr.Image(type="filepath", label="엔딩 프레임 이미지 (도착 화면)")
                            v_reference_images = gr.File(
                                file_count="multiple",
                                label="캐릭터/소재 참조 사진(최대 3장) 및 문서 파일(자동 요약)",
                            )

                vid_inputs = [
                    api_key_vid, video_model, video_prompt, v_first_frame, v_last_frame, v_reference_images,
                    v_aspect_ratio, v_resolution, v_duration, v_auto_extend, v_target_dur, last_video_state,
                ]
                vid_outputs = [video_status, video_output, last_video_state]

                gen_v_evt = generate_vid_btn.click(fn=standard_vid_ui, inputs=vid_inputs, outputs=vid_outputs, api_name=False)
                ext_v_evt = extend_vid_btn.click(fn=extend_vid_ui, inputs=vid_inputs, outputs=vid_outputs, api_name=False)
                stop_vid_btn.click(fn=None, inputs=None, outputs=None, cancels=[gen_v_evt, ext_v_evt], api_name=False)

                def update_video_ui(model_name):
                    # 해상도
                    if "Lite" in model_name:
                        res_up = gr.update(choices=["1080p"], value="1080p")
                    else:
                        res_up = gr.update(choices=["720p", "1080p", "4k"], value="1080p")
                        
                    # 마지막 프레임 및 연장 기능 (Veo 3.0은 미지원)
                    if "3.0" in model_name:
                        last_fr_up = gr.update(interactive=False, value=None, label="엔딩 프레임 (Veo 3.1 이상 지원)")
                        auto_ext_up = gr.update(interactive=False, value=False)
                        extend_btn_up = gr.update(interactive=False)
                    else:
                        last_fr_up = gr.update(interactive=True, label="엔딩 프레임 이미지 (도착 화면)")
                        auto_ext_up = gr.update(interactive=True)
                        extend_btn_up = gr.update(interactive=True)
                        
                    return res_up, last_fr_up, auto_ext_up, extend_btn_up

                video_model.change(
                    fn=update_video_ui, 
                    inputs=[video_model], 
                    outputs=[v_resolution, v_last_frame, v_auto_extend, extend_vid_btn]
                )

                v_auto_extend.change(
                    fn=lambda x: gr.update(interactive=x),
                    inputs=[v_auto_extend],
                    outputs=[v_target_dur]
                )

    return demo


def _resolve_port():
    import socket
    env_port = os.environ.get("APP_PORT") or os.environ.get("GRADIO_SERVER_PORT") or os.environ.get("PORT")
    if env_port:
        try:
            return int(env_port)
        except ValueError:
            pass
    for port in range(7861, 7899):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                s.bind(("0.0.0.0", port))
                return port
            except OSError:
                continue
    return 7860


if __name__ == "__main__":
    from fastapi import FastAPI
    import uvicorn

    demo = build_ui()
    app_api = FastAPI()
    PORT = _resolve_port()

    @app_api.post("/shutdown")
    def shutdown():
        print(f"💡 브라우저 연결 종료가 감지되었습니다. 서버 포트({PORT})를 안전하게 해제하고 종료합니다.")
        os.kill(os.getpid(), signal.SIGINT)
        return {"status": "shutting down"}

    from fastapi.responses import HTMLResponse, FileResponse, PlainTextResponse

    MANUAL_DOCX = os.path.join(BASE_DIR, "사용설명서.docx")
    _manual_cache = {"html": None}

    def _render_manual_html() -> str:
        if _manual_cache["html"] is not None:
            return _manual_cache["html"]
        if not os.path.exists(MANUAL_DOCX):
            return "<h2>사용설명서 파일을 찾을 수 없습니다.</h2>"
        try:
            import mammoth
            with open(MANUAL_DOCX, "rb") as f:
                body = mammoth.convert_to_html(f).value
        except Exception as e:
            try:
                from docx import Document
                doc = Document(MANUAL_DOCX)
                paras = "".join(f"<p>{(p.text or '').replace('<','&lt;').replace('>','&gt;')}</p>" for p in doc.paragraphs)
                body = paras or f"<p>변환 오류: {e}</p>"
            except Exception as e2:
                body = f"<pre>사용설명서 변환 실패: {e2}</pre>"

        page = f"""<!DOCTYPE html>
<html lang=\"ko\"><head><meta charset=\"utf-8\">
<title>📖 Image &amp; Video Laboratory 사용설명서</title>
<style>
  body {{ background:#1a1a1a; color:#eee; font-family:'Pretendard','Apple SD Gothic Neo','맑은 고딕',sans-serif;
         max-width:880px; margin:0 auto; padding:32px 24px 80px; line-height:1.7; }}
  h1, h2, h3 {{ color:#ffd966; border-bottom:1px solid #444; padding-bottom:6px; margin-top:1.6em; }}
  h1 {{ font-size:1.8rem; }} h2 {{ font-size:1.4rem; }} h3 {{ font-size:1.15rem; }}
  p, li {{ font-size:15px; }}
  code, pre {{ background:#2a2a2a; color:#fae; padding:2px 6px; border-radius:4px; }}
  pre {{ padding:12px; overflow-x:auto; }}
  a {{ color:#7ec8ff; }}
  table {{ border-collapse:collapse; margin:10px 0; }}
  td, th {{ border:1px solid #555; padding:6px 10px; }}
  img {{ max-width:100%; height:auto; }}
  .download-btn {{ position:fixed; top:14px; right:18px; background:#ffd966; color:#000;
                  padding:8px 14px; border-radius:8px; font-weight:600; text-decoration:none;
                  box-shadow:0 2px 8px rgba(0,0,0,.4); }}
  .download-btn:hover {{ background:#ffe7a0; }}
</style></head><body>
<a class=\"download-btn\" href=\"manual/download\">⬇️ 원본(.docx) 다운로드</a>
{body}
</body></html>"""
        _manual_cache["html"] = page
        return page

    @app_api.get("/manual", response_class=HTMLResponse)
    def manual_page():
        return HTMLResponse(_render_manual_html())

    @app_api.get("/manual/download")
    def manual_download():
        if not os.path.exists(MANUAL_DOCX):
            return PlainTextResponse("사용설명서 파일이 없습니다.", status_code=404)
        return FileResponse(
            MANUAL_DOCX,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="사용설명서.docx",
        )

    launched_by_parent = bool(
        os.environ.get("LAUNCHED_BY_SCRIPT")
        or os.environ.get("APP_PORT")
        or os.environ.get("GRADIO_SERVER_PORT")
    )

    if launched_by_parent:
        @app_api.on_event("startup")
        async def _signal_ready():
            flag = os.path.join(BASE_DIR, "server_ready.flag")
            try:
                with open(flag, "w") as f:
                    f.write("1")
                print(f"[ready] flag written: {flag}", flush=True)
            except Exception as e:
                print(f"[ready] failed to write flag: {e}", flush=True)
    else:
        def open_browser():
            time.sleep(3)
            webbrowser.open(f"http://127.0.0.1:{PORT}")
        threading.Thread(target=open_browser, daemon=True).start()

    # mysite의 nginx는 /p/<port>/ 로 들어오는 요청의 프리픽스를 떼고 proxy_pass 하므로,
    # Gradio가 자산 URL을 정확히 생성하려면 root_path 를 /p/<port> 로 알려줘야 한다.
    root_path = os.environ.get("GRADIO_ROOT_PATH")
    if root_path is None and launched_by_parent:
        root_path = f"/p/{PORT}"

    print(f"[Image Generator] Starting on port {PORT} (root_path={root_path or '/'})", flush=True)
    app_api = gr.mount_gradio_app(app_api, demo, path="/", root_path=root_path)
    uvicorn.run(app_api, host="0.0.0.0", port=PORT)
